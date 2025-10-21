"""
WORD_DOCUMENTOS_FUSION
"""

#!/usr/bin/env python
# coding: utf-8

# In[1]:


from docx import Document

def combine_word_documents(files, output_filename):
    merged_document = Document()

    for file in files:
        # Cargar el documento Word
        sub_doc = Document(file)

        # Si no es el primer documento, añadir un salto de página
        if sub_doc != files[0]:
            merged_document.add_page_break()

        # Combina el contenido de cada documento
        for element in sub_doc.element.body:
            merged_document.element.body.append(element)

    # Guardar el documento combinado con el nombre especificado
    merged_document.save(output_filename)

# Lista de archivos Word a combinar
file_list = ["C:\\Users\\HP\\Desktop\\CATO CURSOS-1-2024\\GER-TI CATO1-2024\\Cursos\\SEMANA 13\\TEXTOS DE VIDEOS\\PREGUNTAS\\preguntas BOOTSTRAPING.docx",
             "C:\\Users\\HP\\Desktop\\CATO CURSOS-1-2024\\GER-TI CATO1-2024\\Cursos\\SEMANA 13\\TEXTOS DE VIDEOS\\PREGUNTAS\\preguntas CAPITAL RIESGO.docx",
             "C:\\Users\\HP\\Desktop\\CATO CURSOS-1-2024\\GER-TI CATO1-2024\\Cursos\\SEMANA 13\\TEXTOS DE VIDEOS\\PREGUNTAS\\preguntas CROWDFUNDING.docx",
             "C:\\Users\\HP\\Desktop\\CATO CURSOS-1-2024\\GER-TI CATO1-2024\\Cursos\\SEMANA 13\\TEXTOS DE VIDEOS\\PREGUNTAS\\preguntas INCUBADORAS DE EMPRESAS.docx",
             "C:\\Users\\HP\\Desktop\\CATO CURSOS-1-2024\\GER-TI CATO1-2024\\Cursos\\SEMANA 13\\TEXTOS DE VIDEOS\\PREGUNTAS\\preguntas INVERSORES ANGEL.docx"
            ]

output_file = "C:\\Users\\HP\\Desktop\\CATO CURSOS-1-2024\\GER-TI CATO1-2024\\Cursos\\SEMANA 13\\TEXTOS DE VIDEOS\\PREGUNTAS\\full_texto.docx"
# Llamar a la función para combinar los documentos
combine_word_documents(file_list, output_file)


# In[2]:


from docx import Document

def combine_word_documents(files, output_filename):
    merged_document = Document()

    for file in files:
        sub_doc = Document(file)

        # Añadir un salto de página entre los contenidos de cada archivo, excepto antes del primero
        if sub_doc != files[0]:
            merged_document.add_page_break()

        # Copiar cada párrafo del documento actual al documento combinado
        for para in sub_doc.paragraphs:
            merged_document.add_paragraph(para.text, style=para.style.name)

        # Copiar cada tabla del documento actual al documento combinado
        for table in sub_doc.tables:
            # Agrega una nueva tabla con las mismas dimensiones
            new_table = merged_document.add_table(rows=0, cols=len(table.columns))
            new_table.style = table.style
            for row in table.rows:
                cells = new_table.add_row().cells
                for idx, cell in enumerate(row.cells):
                    cells[idx].text = cell.text

    # Guardar el documento combinado con el nombre especificado
    merged_document.save(output_filename)

# Lista de archivos Word a combinar
file_list = ["C:\\Users\\HP\\Desktop\\CATO CURSOS-1-2024\\GER-TI CATO1-2024\\Cursos\\SEMANA 13\\TEXTOS DE VIDEOS\\PREGUNTAS\\preguntas BOOTSTRAPING.docx",
             "C:\\Users\\HP\\Desktop\\CATO CURSOS-1-2024\\GER-TI CATO1-2024\\Cursos\\SEMANA 13\\TEXTOS DE VIDEOS\\PREGUNTAS\\preguntas CAPITAL RIESGO.docx",
             "C:\\Users\\HP\\Desktop\\CATO CURSOS-1-2024\\GER-TI CATO1-2024\\Cursos\\SEMANA 13\\TEXTOS DE VIDEOS\\PREGUNTAS\\preguntas CROWDFUNDING.docx",
             "C:\\Users\\HP\\Desktop\\CATO CURSOS-1-2024\\GER-TI CATO1-2024\\Cursos\\SEMANA 13\\TEXTOS DE VIDEOS\\PREGUNTAS\\preguntas INCUBADORAS DE EMPRESAS.docx",
             "C:\\Users\\HP\\Desktop\\CATO CURSOS-1-2024\\GER-TI CATO1-2024\\Cursos\\SEMANA 13\\TEXTOS DE VIDEOS\\PREGUNTAS\\preguntas INVERSORES ANGEL.docx"
            ]

output_file = "C:\\Users\\HP\\Desktop\\CATO CURSOS-1-2024\\GER-TI CATO1-2024\\Cursos\\SEMANA 13\\TEXTOS DE VIDEOS\\PREGUNTAS\\full_texto.docx"
# Llamar a la función para
# Llamar a la función para combinar los documentos
combine_word_documents(file_list, output_file)


# In[ ]:






if __name__ == "__main__":
    pass
